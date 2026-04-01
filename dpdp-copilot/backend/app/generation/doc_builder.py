"""
DOCX file builder using python-docx.

Converts rendered markdown text into a professional DOCX document
with headers, tables, and formatting.
"""

import os
import re
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger("doc_builder")


def markdown_to_docx(markdown_text: str, output_path: str, title: str = "") -> str:
    """
    Convert rendered markdown text to a DOCX file.

    Simple conversion: handles headers (#, ##, ###), bold (**text**),
    bullet lists (- item), tables (| col |), and paragraphs.

    Args:
        markdown_text: Rendered markdown content
        output_path: Full path to save the DOCX file
        title: Optional document title for the header

    Returns:
        Path to the created DOCX file
    """
    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip horizontal rules
        if line.startswith("---") or line.startswith("***"):
            i += 1
            continue

        # Headers
        if line.startswith("### "):
            _add_heading(doc, line[4:], level=3)
        elif line.startswith("## "):
            _add_heading(doc, line[3:], level=2)
        elif line.startswith("# "):
            _add_heading(doc, line[2:], level=1)

        # Table
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue  # Don't increment i again

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, text)

        # Blockquote
        elif line.startswith("> "):
            text = line[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Regular paragraph
        else:
            if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
                # Italic line
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.italic = True
            else:
                p = doc.add_paragraph()
                _add_formatted_text(p, line)

        i += 1

    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")
    return output_path


def _add_heading(doc, text: str, level: int):
    """Add a heading, stripping markdown formatting."""
    clean = _strip_markdown(text)
    doc.add_heading(clean, level=level)


def _add_formatted_text(paragraph, text: str):
    """Add text with basic markdown formatting (bold, italic)."""
    # Split on bold markers
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_table(doc, table_lines: list[str]):
    """Parse markdown table and add to document."""
    if len(table_lines) < 2:
        return

    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split("|") if cell.strip()]

    # Skip separator line (|---|---|)
    data_start = 1
    if len(table_lines) > 1 and re.match(r"^[\|\s\-:]+$", table_lines[1]):
        data_start = 2

    # Parse data rows
    rows = []
    for line in table_lines[data_start:]:
        cells = [cell.strip() for cell in line.split("|") if cell.strip()]
        if cells:
            rows.append(cells)

    if not headers:
        return

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = _strip_markdown(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(headers):
                table.rows[i + 1].cells[j].text = _strip_markdown(cell_text)

    doc.add_paragraph()  # Space after table


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting characters."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # Bold
    text = re.sub(r"\*(.*?)\*", r"\1", text)  # Italic
    text = re.sub(r"`(.*?)`", r"\1", text)  # Code
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)  # Links
    return text.strip()
